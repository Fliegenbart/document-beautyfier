from __future__ import annotations

from datetime import date
from dataclasses import dataclass
from pathlib import Path
from typing import Iterator

from docx import Document
from docx.document import Document as DocxDocumentType
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Cm, Pt, RGBColor
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from reportlab.lib import colors
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.lib.utils import ImageReader
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


@dataclass
class TemplateConfig:
    margin_cm: float
    body_size: float
    h1_size: float
    h2_size: float
    h3_size: float
    title_size: float
    line_spacing: float
    header_line_size: str
    logo_width_cm: float


@dataclass
class PdfThemeConfig:
    default_font: str
    cover_tagline: str
    default_primary_hex: str
    default_text_hex: str
    summary_fill_hex: str
    table_row_fill_a: str
    table_row_fill_b: str
    footer_color_hex: str
    cover_rule_height_cm: float


TEMPLATES = {
    "minimal": TemplateConfig(2.4, 10.8, 19, 14, 12.5, 25, 1.13, "8", 4.8),
    "executive": TemplateConfig(2.1, 11, 20, 15, 13, 28, 1.15, "12", 5.5),
    "bold": TemplateConfig(1.8, 11.2, 22, 16, 14, 30, 1.2, "16", 6.0),
}

PDF_THEMES = {
    "consulting": PdfThemeConfig(
        default_font="Helvetica",
        cover_tagline="Business Strategy Whitepaper",
        default_primary_hex="D7263D",
        default_text_hex="1A1A1A",
        summary_fill_hex="F6F8FB",
        table_row_fill_a="FFFFFF",
        table_row_fill_b="F6F8FA",
        footer_color_hex="666666",
        cover_rule_height_cm=0.24,
    ),
    "technical": PdfThemeConfig(
        default_font="Helvetica",
        cover_tagline="Technical Whitepaper",
        default_primary_hex="005A9C",
        default_text_hex="132B43",
        summary_fill_hex="F3F7FC",
        table_row_fill_a="F9FBFE",
        table_row_fill_b="EDF4FB",
        footer_color_hex="4F5E72",
        cover_rule_height_cm=0.2,
    ),
    "regulatory": PdfThemeConfig(
        default_font="Times-Roman",
        cover_tagline="Regulatory Assessment Whitepaper",
        default_primary_hex="6E4C2C",
        default_text_hex="2A221C",
        summary_fill_hex="FAF7F2",
        table_row_fill_a="FFFEFC",
        table_row_fill_b="F7F1E9",
        footer_color_hex="5E5348",
        cover_rule_height_cm=0.22,
    ),
}


def normalize_color_string(value: str) -> str:
    text = value.strip().lower().replace(" ", "")
    if text.startswith("#"):
        text = text[1:]
    if "," in text:
        parts = text.split(",")
        if len(parts) != 3:
            raise ValueError(f"Invalid RGB color: {value}")
        rgb = [max(0, min(255, int(p))) for p in parts]
        return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    if len(text) == 6 and all(c in "0123456789abcdef" for c in text):
        return text.upper()
    raise ValueError(f"Invalid color: {value} (use #F50000 or 245,0,0)")


def hex_to_rgb(hex_color: str) -> RGBColor:
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color_hex: str, size: str) -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)


def clear_header_paragraphs(header) -> None:
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)


def style_named_styles(doc: Document, font_name: str, text_rgb: RGBColor, primary_rgb: RGBColor, cfg: TemplateConfig) -> None:
    normal = doc.styles["Normal"]
    normal_font = normal.font
    normal_font.name = font_name
    normal_font.size = Pt(cfg.body_size)
    normal_font.color.rgb = text_rgb

    heading_specs = {
        "Heading 1": (Pt(cfg.h1_size), True),
        "Heading 2": (Pt(cfg.h2_size), True),
        "Heading 3": (Pt(cfg.h3_size), True),
    }

    for style_name, (size, bold) in heading_specs.items():
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st_font = st.font
            st_font.name = font_name
            st_font.size = size
            st_font.bold = bold
            st_font.color.rgb = primary_rgb

    if "Title" in doc.styles:
        title = doc.styles["Title"]
        title_font = title.font
        title_font.name = font_name
        title_font.size = Pt(cfg.title_size)
        title_font.bold = True
        title_font.color.rgb = primary_rgb


def style_paragraphs(doc: Document, font_name: str, text_rgb: RGBColor, primary_rgb: RGBColor, cfg: TemplateConfig) -> None:
    first_content_paragraph = None

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        pf = p.paragraph_format
        pf.space_after = Pt(8)
        pf.space_before = Pt(0)
        pf.line_spacing = cfg.line_spacing

        if not text:
            continue

        if first_content_paragraph is None:
            first_content_paragraph = p

        for r in p.runs:
            if not r.font.name:
                r.font.name = font_name
            if not r.font.size:
                r.font.size = Pt(cfg.body_size)
            if r.font.color is None or r.font.color.rgb is None:
                r.font.color.rgb = text_rgb

        style_name = p.style.name if p.style else ""
        if style_name.startswith("Heading"):
            for r in p.runs:
                r.font.color.rgb = primary_rgb

    if first_content_paragraph is not None:
        style_name = first_content_paragraph.style.name if first_content_paragraph.style else ""
        if style_name not in {"Title", "Heading 1"}:
            first_content_paragraph.style = doc.styles["Title"] if "Title" in doc.styles else doc.styles["Heading 1"]
            first_content_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def style_tables(doc: Document, font_name: str, text_rgb: RGBColor, primary_hex: str) -> None:
    white = RGBColor(255, 255, 255)
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        if not table.rows:
            continue

        header_row = table.rows[0]
        for cell in header_row.cells:
            set_cell_shading(cell, primary_hex)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = white
                    r.font.name = font_name
                    r.font.size = Pt(10.5)

        for row in table.rows[1:]:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.font.color is None or r.font.color.rgb is None:
                            r.font.color.rgb = text_rgb
                        if not r.font.name:
                            r.font.name = font_name
                        if not r.font.size:
                            r.font.size = Pt(10.5)


def style_sections_and_header(
    doc: Document,
    logo_path: Path | None,
    org_name: str,
    font_name: str,
    text_rgb: RGBColor,
    primary_hex: str,
    cfg: TemplateConfig,
) -> None:
    for section in doc.sections:
        section.top_margin = Cm(cfg.margin_cm)
        section.bottom_margin = Cm(cfg.margin_cm)
        section.left_margin = Cm(cfg.margin_cm)
        section.right_margin = Cm(cfg.margin_cm)

        section.different_first_page_header_footer = True

        for header in [section.header, section.first_page_header]:
            clear_header_paragraphs(header)
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)

            if logo_path and logo_path.exists():
                run = p.add_run()
                run.add_picture(str(logo_path), width=Cm(cfg.logo_width_cm))
            else:
                run = p.add_run(org_name)
                run.font.name = font_name
                run.font.size = Pt(cfg.body_size)
                run.font.bold = True
                run.font.color.rgb = text_rgb

            set_paragraph_bottom_border(p, color_hex=primary_hex, size=cfg.header_line_size)


def ensure_custom_styles(doc: Document, font_name: str, primary_rgb: RGBColor, body_size: float) -> None:
    if "Beautifier Accent" not in [s.name for s in doc.styles]:
        accent = doc.styles.add_style("Beautifier Accent", WD_STYLE_TYPE.PARAGRAPH)
        accent.base_style = doc.styles["Normal"]
        accent.font.name = font_name
        accent.font.size = Pt(body_size)
        accent.font.bold = True
        accent.font.color.rgb = primary_rgb


def apply_style(
    input_docx: Path,
    output_docx: Path,
    logo: Path | None = None,
    org_name: str = "Your Organization",
    font: str = "Calibri",
    template: str = "executive",
    primary_color: str = "#F50000",
    text_color: str = "#111111",
    line_spacing: float = 1.5,
) -> None:
    if template not in TEMPLATES:
        raise ValueError(f"Unsupported template: {template}")

    primary_hex = normalize_color_string(primary_color)
    text_hex = normalize_color_string(text_color)
    primary_rgb = hex_to_rgb(primary_hex)
    text_rgb = hex_to_rgb(text_hex)
    cfg = TEMPLATES[template]

    doc = Document(str(input_docx))

    style_named_styles(doc, font, text_rgb, primary_rgb, cfg)
    ensure_custom_styles(doc, font, primary_rgb, cfg.body_size)
    cfg_for_docx = TemplateConfig(
        margin_cm=cfg.margin_cm,
        body_size=cfg.body_size,
        h1_size=cfg.h1_size,
        h2_size=cfg.h2_size,
        h3_size=cfg.h3_size,
        title_size=cfg.title_size,
        line_spacing=max(cfg.line_spacing, line_spacing),
        header_line_size=cfg.header_line_size,
        logo_width_cm=cfg.logo_width_cm,
    )
    style_paragraphs(doc, font, text_rgb, primary_rgb, cfg_for_docx)
    style_tables(doc, font, text_rgb, primary_hex)
    style_sections_and_header(doc, logo, org_name, font, text_rgb, primary_hex, cfg)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))


def _iter_block_items(parent: DocxDocumentType) -> Iterator[DocxParagraph | DocxTable]:
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, parent)


def _extract_doc_title(doc: Document) -> str:
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            return text
    return "Styled Document"


def _build_pdf_styles(
    font: str,
    primary_hex: str,
    text_hex: str,
    cfg: TemplateConfig,
    line_spacing: float,
    pdf_theme: PdfThemeConfig,
):
    primary = HexColor(f"#{primary_hex}")
    text = HexColor(f"#{text_hex}")
    sheet = getSampleStyleSheet()
    return {
        "h1": ParagraphStyle(
            "h1",
            parent=sheet["Heading1"],
            fontName=font,
            fontSize=cfg.h1_size,
            leading=cfg.h1_size + 4,
            textColor=primary,
            spaceBefore=8,
            spaceAfter=8,
        ),
        "h2": ParagraphStyle(
            "h2",
            parent=sheet["Heading2"],
            fontName=font,
            fontSize=cfg.h2_size,
            leading=cfg.h2_size + 3,
            textColor=primary,
            spaceBefore=7,
            spaceAfter=6,
        ),
        "h3": ParagraphStyle(
            "h3",
            parent=sheet["Heading3"],
            fontName=font,
            fontSize=cfg.h3_size,
            leading=cfg.h3_size + 3,
            textColor=primary,
            spaceBefore=6,
            spaceAfter=5,
        ),
        "body": ParagraphStyle(
            "body",
            parent=sheet["BodyText"],
            fontName=font,
            fontSize=cfg.body_size,
            leading=cfg.body_size * line_spacing,
            textColor=text,
            spaceAfter=max(6, cfg.body_size * 0.72),
        ),
        "cover_title": ParagraphStyle(
            "cover_title",
            parent=sheet["Title"],
            fontName=font,
            fontSize=cfg.title_size + 4,
            leading=cfg.title_size + 8,
            textColor=primary,
            alignment=1,
            spaceAfter=12,
        ),
        "cover_sub": ParagraphStyle(
            "cover_sub",
            parent=sheet["Normal"],
            fontName=font,
            fontSize=cfg.body_size + 1,
            textColor=text,
            alignment=1,
        ),
        "cover_tagline": ParagraphStyle(
            "cover_tagline",
            parent=sheet["Normal"],
            fontName=font,
            fontSize=cfg.body_size - 0.3,
            textColor=colors.HexColor("#777777"),
            alignment=1,
            spaceAfter=8,
        ),
        "footer": ParagraphStyle(
            "footer",
            parent=sheet["Normal"],
            fontName=font,
            fontSize=8.5,
            textColor=colors.HexColor(f"#{pdf_theme.footer_color_hex}"),
        ),
    }


def _heading_list(doc: Document) -> list[str]:
    headings: list[str] = []
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style_name = (p.style.name or "").lower()
        if "heading 1" in style_name or "heading 2" in style_name:
            headings.append(text)
    return headings[:14]


def _summary_text(doc: Document) -> str:
    chunks: list[str] = []
    for p in doc.paragraphs:
        text = " ".join((p.text or "").split()).strip()
        if not text:
            continue
        style_name = (p.style.name or "").lower()
        if "heading" in style_name:
            continue
        chunks.append(text)
        if len(" ".join(chunks)) > 520:
            break
    if not chunks:
        return "This whitepaper outlines the business challenge, supporting context, and a practical solution path."
    return " ".join(chunks)[:650]


def _reading_width_points(body_size_pt: float, reading_width_ch: int) -> float:
    # Approximate average character width in points for business sans-serif fonts.
    return body_size_pt * reading_width_ch * 0.52


def _logo_image_flowable(logo: Path, max_width_cm: float, max_height_cm: float, align: str = "CENTER") -> Image:
    img_reader = ImageReader(str(logo))
    px_w, px_h = img_reader.getSize()
    if px_w <= 0 or px_h <= 0:
        return Image(str(logo), width=max_width_cm * cm, hAlign=align)

    width_pt = max_width_cm * cm
    height_pt = width_pt * (px_h / px_w)
    max_height_pt = max_height_cm * cm

    if height_pt > max_height_pt:
        scale = max_height_pt / height_pt
        width_pt *= scale
        height_pt = max_height_pt

    return Image(str(logo), width=width_pt, height=height_pt, hAlign=align)


def _render_cover(story, title: str, org_name: str, logo: Path | None, rl_doc, cfg: TemplateConfig, theme: PdfThemeConfig, styles, primary):
    if theme.cover_tagline.startswith("Business"):
        # Consulting: centered, clean with strong accent rule.
        story.append(Spacer(1, 1.2 * cm))
        if logo and logo.exists():
            story.append(_logo_image_flowable(logo, max_width_cm=cfg.logo_width_cm, max_height_cm=2.8, align="CENTER"))
            story.append(Spacer(1, 0.8 * cm))
        story.append(Paragraph(theme.cover_tagline, styles["cover_tagline"]))
        story.append(Paragraph(title, styles["cover_title"]))
        story.append(Paragraph(org_name, styles["cover_sub"]))
        story.append(Paragraph(date.today().strftime("%B %d, %Y"), styles["cover_sub"]))
        story.append(Spacer(1, 0.9 * cm))
        story.append(
            Table(
                [[" "]],
                colWidths=[rl_doc.width],
                rowHeights=[theme.cover_rule_height_cm * cm],
                style=[("BACKGROUND", (0, 0), (-1, -1), primary), ("LINEBELOW", (0, 0), (-1, -1), 0, primary)],
            )
        )
        story.append(PageBreak())
        return

    if theme.cover_tagline.startswith("Technical"):
        # Technical: left-aligned hero with compact metadata.
        story.append(Spacer(1, 0.9 * cm))
        if logo and logo.exists():
            story.append(
                _logo_image_flowable(logo, max_width_cm=(cfg.logo_width_cm - 1.0), max_height_cm=2.4, align="LEFT")
            )
            story.append(Spacer(1, 0.6 * cm))
        story.append(Paragraph(theme.cover_tagline, styles["h3"]))
        story.append(Paragraph(title, styles["cover_title"]))
        meta = Table(
            [
                ["Organization", org_name],
                ["Date", date.today().strftime("%B %d, %Y")],
                ["Profile", "Technical Documentation"],
            ],
            colWidths=[rl_doc.width * 0.25, rl_doc.width * 0.75],
            style=[
                ("BACKGROUND", (0, 0), (0, -1), primary),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.white),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#B5C2CF")),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ],
        )
        story.append(Spacer(1, 0.4 * cm))
        story.append(meta)
        story.append(Spacer(1, 0.6 * cm))
        story.append(
            Table(
                [[" "]],
                colWidths=[rl_doc.width],
                rowHeights=[theme.cover_rule_height_cm * cm],
                style=[("BACKGROUND", (0, 0), (-1, -1), primary)],
            )
        )
        story.append(PageBreak())
        return

    # Regulatory: formal title block.
    story.append(Spacer(1, 1.0 * cm))
    story.append(
        Table(
            [[Paragraph("Regulatory Whitepaper", styles["cover_tagline"])]],
            colWidths=[rl_doc.width],
            style=[
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F3EEE6")),
                ("BOX", (0, 0), (-1, -1), 0.8, primary),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ],
        )
    )
    story.append(Spacer(1, 0.65 * cm))
    story.append(Paragraph(title, styles["cover_title"]))
    story.append(Paragraph(org_name, styles["cover_sub"]))
    story.append(Paragraph(date.today().strftime("%B %d, %Y"), styles["cover_sub"]))
    if logo and logo.exists():
        story.append(Spacer(1, 0.5 * cm))
        story.append(_logo_image_flowable(logo, max_width_cm=(cfg.logo_width_cm - 0.4), max_height_cm=2.6, align="CENTER"))
    story.append(Spacer(1, 0.8 * cm))
    story.append(
        Table(
            [["Confidential - Prepared for internal quality and compliance stakeholders"]],
            colWidths=[rl_doc.width],
            style=[
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FBF8F3")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#65584B")),
                ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#D7CCBF")),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                ("FONTSIZE", (0, 0), (-1, -1), 8.5),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ],
        )
    )
    story.append(PageBreak())


def apply_style_pdf(
    input_docx: Path,
    output_pdf: Path,
    logo: Path | None = None,
    org_name: str = "Your Organization",
    font: str = "auto",
    pdf_theme: str = "consulting",
    template: str = "executive",
    primary_color: str = "#F50000",
    text_color: str = "#111111",
    reading_width_ch: int = 72,
    line_spacing: float = 1.55,
    include_summary_page: bool = True,
) -> None:
    if template not in TEMPLATES:
        raise ValueError(f"Unsupported template: {template}")
    if pdf_theme not in PDF_THEMES:
        raise ValueError(f"Unsupported pdf_theme: {pdf_theme}")

    theme = PDF_THEMES[pdf_theme]
    cfg = TEMPLATES[template]
    primary_hex = theme.default_primary_hex
    text_hex = theme.default_text_hex
    if primary_color and primary_color.strip() and primary_color.strip().lower() not in {"", "auto"}:
        primary_hex = normalize_color_string(primary_color)
    if text_color and text_color.strip() and text_color.strip().lower() not in {"", "auto"}:
        text_hex = normalize_color_string(text_color)

    doc = Document(str(input_docx))
    title = _extract_doc_title(doc)
    if reading_width_ch < 50 or reading_width_ch > 80:
        raise ValueError("reading_width_ch must be between 50 and 80.")
    if line_spacing < 1.4 or line_spacing > 2.0:
        raise ValueError("line_spacing should be between 1.4 and 2.0.")

    selected_font = theme.default_font if (not font or font.strip().lower() == "auto") else font
    styles = _build_pdf_styles(selected_font, primary_hex, text_hex, cfg, line_spacing, theme)

    output_pdf.parent.mkdir(parents=True, exist_ok=True)

    base_margin = cfg.margin_cm * cm
    frame_max_width = _reading_width_points(cfg.body_size, reading_width_ch)
    usable_width = A4[0] - (2 * base_margin)
    frame_width = min(usable_width, frame_max_width)
    side_margin = (A4[0] - frame_width) / 2

    rl_doc = SimpleDocTemplate(
        str(output_pdf),
        pagesize=A4,
        leftMargin=side_margin,
        rightMargin=side_margin,
        topMargin=base_margin,
        bottomMargin=base_margin,
    )

    primary = HexColor(f"#{primary_hex}")
    story = []

    _render_cover(story, title, org_name, logo, rl_doc, cfg, theme, styles, primary)

    if include_summary_page:
        summary = _summary_text(doc)
        heads = _heading_list(doc)
        story.append(Paragraph("Executive Summary", styles["h1"]))
        summary_box = Table(
            [[Paragraph(summary, styles["body"])]],
            colWidths=[rl_doc.width],
            style=[
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(f"#{theme.summary_fill_hex}")),
                ("BOX", (0, 0), (-1, -1), 0.8, primary),
                ("LEFTPADDING", (0, 0), (-1, -1), 10),
                ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ],
        )
        story.append(summary_box)
        story.append(Spacer(1, 0.5 * cm))
        if heads:
            story.append(Paragraph("Contents", styles["h2"]))
            for heading in heads:
                story.append(Paragraph(f"- {heading}", styles["body"]))
        story.append(PageBreak())

    # Content
    for item in _iter_block_items(doc):
        if isinstance(item, DocxParagraph):
            text = (item.text or "").strip()
            if not text:
                continue
            style_name = (item.style.name or "").lower()
            if "heading 1" in style_name:
                story.append(Paragraph(text, styles["h1"]))
            elif "heading 2" in style_name:
                story.append(Paragraph(text, styles["h2"]))
            elif "heading 3" in style_name:
                story.append(Paragraph(text, styles["h3"]))
            else:
                story.append(Paragraph(text.replace("\n", "<br/>"), styles["body"]))
        else:
            rows = []
            for row in item.rows:
                rows.append([" ".join(cell.text.split()) for cell in row.cells])
            if not rows:
                continue
            col_count = max(len(r) for r in rows)
            table = Table(rows, colWidths=[rl_doc.width / col_count] * col_count, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), primary),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, -1), selected_font),
                        ("FONTSIZE", (0, 0), (-1, 0), 10),
                        ("FONTSIZE", (0, 1), (-1, -1), 9),
                        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#BBBBBB")),
                        ("ALIGN", (0, 0), (-1, 0), "LEFT"),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        (
                            "ROWBACKGROUNDS",
                            (0, 1),
                            (-1, -1),
                            [colors.HexColor(f"#{theme.table_row_fill_a}"), colors.HexColor(f"#{theme.table_row_fill_b}")],
                        ),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ]
                )
            )
            story.append(Spacer(1, 0.2 * cm))
            story.append(table)
            story.append(Spacer(1, 0.4 * cm))

    def draw_frame(canvas, _doc):
        canvas.saveState()
        canvas.setStrokeColor(primary)
        canvas.setLineWidth(1)
        canvas.line(_doc.leftMargin, A4[1] - _doc.topMargin + 8, A4[0] - _doc.rightMargin, A4[1] - _doc.topMargin + 8)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor(f"#{theme.footer_color_hex}"))
        canvas.drawString(_doc.leftMargin, _doc.bottomMargin - 16, org_name)
        canvas.drawRightString(A4[0] - _doc.rightMargin, _doc.bottomMargin - 16, f"Page {canvas.getPageNumber()}")
        canvas.restoreState()

    rl_doc.build(story, onFirstPage=draw_frame, onLaterPages=draw_frame)
