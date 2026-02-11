#!/usr/bin/env python3
from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


@dataclass
class TemplateConfig:
    margin_cm: float
    body_size: float
    h1_size: float
    h2_size: float
    h3_size: float
    title_size: float
    line_spacing: float
    header_line_size: str
    logo_width_cm: float


TEMPLATES = {
    "minimal": TemplateConfig(2.4, 10.8, 19, 14, 12.5, 25, 1.13, "8", 4.8),
    "executive": TemplateConfig(2.1, 11, 20, 15, 13, 28, 1.15, "12", 5.5),
    "bold": TemplateConfig(1.8, 11.2, 22, 16, 14, 30, 1.2, "16", 6.0),
}


def normalize_color_string(value: str) -> str:
    text = value.strip().lower().replace(" ", "")
    if text.startswith("#"):
        text = text[1:]
    if "," in text:
        parts = text.split(",")
        if len(parts) != 3:
            raise ValueError(f"Ungueltige RGB-Farbe: {value}")
        rgb = [max(0, min(255, int(p))) for p in parts]
        return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    if len(text) == 6 and all(c in "0123456789abcdef" for c in text):
        return text.upper()
    raise ValueError(f"Ungueltige Farbe: {value} (nutze z.B. #F50000 oder 245,0,0)")


def hex_to_rgb(hex_color: str) -> RGBColor:
    return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))


def set_cell_shading(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_paragraph_bottom_border(paragraph, color_hex: str, size: str) -> None:
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)

    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)

    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)


def clear_header_paragraphs(header) -> None:
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)


def style_named_styles(doc: Document, font_name: str, text_rgb: RGBColor, primary_rgb: RGBColor, cfg: TemplateConfig) -> None:
    normal = doc.styles["Normal"]
    normal_font = normal.font
    normal_font.name = font_name
    normal_font.size = Pt(cfg.body_size)
    normal_font.color.rgb = text_rgb

    heading_specs = {
        "Heading 1": (Pt(cfg.h1_size), True),
        "Heading 2": (Pt(cfg.h2_size), True),
        "Heading 3": (Pt(cfg.h3_size), True),
    }

    for style_name, (size, bold) in heading_specs.items():
        if style_name in doc.styles:
            st = doc.styles[style_name]
            st_font = st.font
            st_font.name = font_name
            st_font.size = size
            st_font.bold = bold
            st_font.color.rgb = primary_rgb

    if "Title" in doc.styles:
        title = doc.styles["Title"]
        title_font = title.font
        title_font.name = font_name
        title_font.size = Pt(cfg.title_size)
        title_font.bold = True
        title_font.color.rgb = primary_rgb


def style_paragraphs(doc: Document, font_name: str, text_rgb: RGBColor, primary_rgb: RGBColor, cfg: TemplateConfig) -> None:
    first_content_paragraph = None

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        pf = p.paragraph_format
        pf.space_after = Pt(8)
        pf.space_before = Pt(0)
        pf.line_spacing = cfg.line_spacing

        if not text:
            continue

        if first_content_paragraph is None:
            first_content_paragraph = p

        for r in p.runs:
            if not r.font.name:
                r.font.name = font_name
            if not r.font.size:
                r.font.size = Pt(cfg.body_size)
            if r.font.color is None or r.font.color.rgb is None:
                r.font.color.rgb = text_rgb

        style_name = p.style.name if p.style else ""
        if style_name.startswith("Heading"):
            for r in p.runs:
                r.font.color.rgb = primary_rgb

    if first_content_paragraph is not None:
        style_name = first_content_paragraph.style.name if first_content_paragraph.style else ""
        if style_name not in {"Title", "Heading 1"}:
            first_content_paragraph.style = doc.styles["Title"] if "Title" in doc.styles else doc.styles["Heading 1"]
            first_content_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def style_tables(doc: Document, font_name: str, text_rgb: RGBColor, primary_hex: str) -> None:
    white = RGBColor(255, 255, 255)
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        if not table.rows:
            continue

        header_row = table.rows[0]
        for cell in header_row.cells:
            set_cell_shading(cell, primary_hex)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = white
                    r.font.name = font_name
                    r.font.size = Pt(10.5)

        for row in table.rows[1:]:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        if r.font.color is None or r.font.color.rgb is None:
                            r.font.color.rgb = text_rgb
                        if not r.font.name:
                            r.font.name = font_name
                        if not r.font.size:
                            r.font.size = Pt(10.5)


def style_sections_and_header(
    doc: Document,
    logo_path: Path | None,
    org_name: str,
    font_name: str,
    text_rgb: RGBColor,
    primary_hex: str,
    cfg: TemplateConfig,
) -> None:
    for section in doc.sections:
        section.top_margin = Cm(cfg.margin_cm)
        section.bottom_margin = Cm(cfg.margin_cm)
        section.left_margin = Cm(cfg.margin_cm)
        section.right_margin = Cm(cfg.margin_cm)

        section.different_first_page_header_footer = True

        for header in [section.header, section.first_page_header]:
            clear_header_paragraphs(header)
            p = header.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(6)

            if logo_path and logo_path.exists():
                run = p.add_run()
                run.add_picture(str(logo_path), width=Cm(cfg.logo_width_cm))
            else:
                run = p.add_run(org_name)
                run.font.name = font_name
                run.font.size = Pt(cfg.body_size)
                run.font.bold = True
                run.font.color.rgb = text_rgb

            set_paragraph_bottom_border(p, color_hex=primary_hex, size=cfg.header_line_size)


def ensure_custom_styles(doc: Document, font_name: str, primary_rgb: RGBColor, body_size: float) -> None:
    if "Beautifier Accent" not in [s.name for s in doc.styles]:
        accent = doc.styles.add_style("Beautifier Accent", WD_STYLE_TYPE.PARAGRAPH)
        accent.base_style = doc.styles["Normal"]
        accent.font.name = font_name
        accent.font.size = Pt(body_size)
        accent.font.bold = True
        accent.font.color.rgb = primary_rgb


def main() -> None:
    parser = argparse.ArgumentParser(description="Apply configurable business styling to DOCX whitepapers.")
    parser.add_argument("input_docx", type=Path)
    parser.add_argument("output_docx", type=Path)
    parser.add_argument("--logo", type=Path, default=None, help="Path to logo image (png/jpg).")
    parser.add_argument("--org-name", type=str, default="Your Organization")
    parser.add_argument("--font", type=str, default="Calibri")
    parser.add_argument("--template", choices=sorted(TEMPLATES.keys()), default="executive")
    parser.add_argument("--primary-color", type=str, default="#F50000", help="Hex or RGB string, e.g. #F50000 or 245,0,0")
    parser.add_argument("--text-color", type=str, default="#111111", help="Hex or RGB string, e.g. #111111 or 17,17,17")
    args = parser.parse_args()

    primary_hex = normalize_color_string(args.primary_color)
    text_hex = normalize_color_string(args.text_color)
    primary_rgb = hex_to_rgb(primary_hex)
    text_rgb = hex_to_rgb(text_hex)

    cfg = TEMPLATES[args.template]

    doc = Document(str(args.input_docx))

    style_named_styles(doc, args.font, text_rgb, primary_rgb, cfg)
    ensure_custom_styles(doc, args.font, primary_rgb, cfg.body_size)
    style_paragraphs(doc, args.font, text_rgb, primary_rgb, cfg)
    style_tables(doc, args.font, text_rgb, primary_hex)
    style_sections_and_header(doc, args.logo, args.org_name, args.font, text_rgb, primary_hex, cfg)

    args.output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(args.output_docx))


if __name__ == "__main__":
    main()
